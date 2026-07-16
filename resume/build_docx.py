#!/usr/bin/env python3
"""Rebuild the downloadable resume DOCX from resume.md.

Generates an ATS-safe DOCX (single column, standard headings, plain bullets,
no tables, no text boxes, no images) and writes it straight to the site's
download asset.

Usage:
    pip install python-docx
    python3 build_docx.py
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "resume.md")
OUT = os.path.join(HERE, "..", "assets", "Jonathan-Brenes-Fernandez-Resume.docx")

doc = Document()

# Base styling: single, clean, ATS-friendly font.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

for section in doc.sections:
    section.left_margin = section.right_margin = Pt(54)   # ~0.75in
    section.top_margin = section.bottom_margin = Pt(54)

lines = open(SRC, encoding="utf-8").read().splitlines()


def add_inline(paragraph, text):
    """Render **bold** segments; strip other markdown."""
    for seg in re.split(r"(\*\*.*?\*\*)", text):
        if not seg:
            continue
        run = paragraph.add_run(seg[2:-2] if seg.startswith("**") else seg)
        if seg.startswith("**"):
            run.bold = True


for raw in lines:
    line = raw.rstrip()
    if not line:
        continue
    if line.startswith("---"):
        continue

    if line.startswith("# "):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line[2:].strip())
        run.bold = True
        run.font.size = Pt(20)
    elif line.startswith("## "):
        p = doc.add_paragraph()
        p.space_before = Pt(10)
        run = p.add_run(line[3:].strip().upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x57)
    elif line.startswith("### "):
        p = doc.add_paragraph()
        run = p.add_run(line[4:].strip())
        run.bold = True
        run.font.size = Pt(11)
    elif line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, line[2:].strip())
    else:
        p = doc.add_paragraph()
        add_inline(p, line.strip())
        # Center the contact line (line right under the name).
        if "@" in line and "|" in line:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

out = os.path.normpath(OUT)
doc.save(out)
print("wrote", out)
